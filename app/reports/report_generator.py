import io
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from app.ai.factory import AIFactory
from app.reports.prompts import (
    SYSTEM_INSTRUCTION_RSSI, 
    PROMPT_EXECUTIVE_SUMMARY, 
    PROMPT_VULN_DETAILS, 
    PROMPT_CONCLUSION
)

class ReportGenerator:
    def __init__(self):
        try:
            self.ai_provider = AIFactory.create_provider()
        except Exception as e:
            print(f"AI Provider Init Failed: {e}")
            self.ai_provider = None

    def generate_report(self, scan):
        """
        Main orchestration method.
        Returns a BytesIO object containing the Docx file.
        """
        # 1. Prepare Data
        stats = self._calculate_stats(scan)
        grouped_vulns = self._group_vulnerabilities(scan)
        top_risks_text = self._format_top_risks(grouped_vulns)
        
        # Get Language & Company Config
        from app.models_settings import Settings
        settings = Settings.query.first()
        
        lang_code = settings.language if settings and settings.language else 'fr'
        lang_map = {'fr': 'French', 'en': 'English', 'es': 'Spanish'}
        target_lang = lang_map.get(lang_code, 'French')
        
        company_name = settings.company_name if settings and settings.company_name else "Spectra"
        
        # Prepare System Instruction
        system_prompt = SYSTEM_INSTRUCTION_RSSI.format(
            language=target_lang,
            company_name=company_name
        )
        
        # 2. AI Generation - Step 1: Executive Summary
        exec_summary = self._ai_generate_summary(scan.project.name, stats, top_risks_text, target_lang, system_prompt)
        
        # 3. AI Generation - Step 2: Critical/High Details
        detailed_analyses = {}
        priority_groups = {k: v for k, v in grouped_vulns.items() if v['severity'] in ['CRITICAL', 'HIGH']}
        
        for signature, data in priority_groups.items():
            analysis = self._ai_generate_vuln_details(data, target_lang, system_prompt)
            detailed_analyses[signature] = analysis

        # 4. AI Generation - Step 3: Conclusion
        conclusion = self._ai_generate_conclusion(stats, target_lang, system_prompt)
        
        # 5. Build Word Document
        return self._build_docx(scan, stats, grouped_vulns, exec_summary, detailed_analyses, conclusion, target_lang)

    def _calculate_stats(self, scan):
        results = scan.results
        return {
            'total': len(results),
            'critical': len([r for r in results if r.severity == 'CRITICAL']),
            'high': len([r for r in results if r.severity == 'HIGH']),
            'medium': len([r for r in results if r.severity == 'MEDIUM']),
            'low': len([r for r in results if r.severity in ['LOW', 'INFO', 'UNKNOWN']])
        }

    def _group_vulnerabilities(self, scan):
        groups = {}
        for vuln in scan.results:
            signature = f"{vuln.vuln_id}" 
            
            if signature not in groups:
                groups[signature] = {
                    'title': vuln.title,
                    'tool': vuln.tool,
                    'severity': vuln.severity,
                    'owasp': vuln.owasp_category,
                    'vuln_id': vuln.vuln_id,
                    'description': vuln.description, 
                    'locations': []
                }
            
            groups[signature]['locations'].append({
                'path': vuln.file_path,
                'line': vuln.line_number
            })
            
        return groups

    def _format_top_risks(self, grouped_vulns):
        weights = {'CRITICAL': 4, 'HIGH': 3, 'MEDIUM': 2, 'LOW': 1, 'INFO': 0}
        
        sorted_groups = sorted(
            grouped_vulns.values(),
            key=lambda x: weights.get(x['severity'], 0),
            reverse=True
        )
        
        top_3 = sorted_groups[:3]
        
        text = ""
        for i, item in enumerate(top_3, 1):
            text += f"{i}. [{item['severity']}] {item['title']} ({item['tool']})\n"
            
        if not text:
            text = "No major vulnerabilities detected."
            
        return text

    def _ai_generate_summary(self, project_name, stats, top_risks_text, language, system_instruction):
        if not self.ai_provider:
            return "AI Provider not configured. Executive summary unavailable."
            
        prompt = PROMPT_EXECUTIVE_SUMMARY.format(
            project_name=project_name,
            total_count=stats['total'],
            critical_count=stats['critical'],
            high_count=stats['high'],
            medium_count=stats['medium'],
            low_count=stats['low'],
            top_3_risks_text=top_risks_text,
            language=language
        )
        
        return self.ai_provider.generate(prompt, system_instruction=system_instruction)

    def _ai_generate_vuln_details(self, data, language, system_instruction):
        if not self.ai_provider:
            return "Description unavailable (No AI)."
            
        prompt = PROMPT_VULN_DETAILS.format(
            title=data['title'],
            owasp_category=data['owasp'],
            tool=data['tool'],
            severity=data['severity'],
            description=data['description'],
            language=language
        )
        
        return self.ai_provider.generate(prompt, system_instruction=system_instruction)

    def _ai_generate_conclusion(self, stats, language, system_instruction):
        if not self.ai_provider:
            return "Conclusion unavailable."
            
        prompt = PROMPT_CONCLUSION.format(
            total_count=stats['total'],
            critical_count=stats['critical'],
            language=language
        )
        
        return self.ai_provider.generate(prompt, system_instruction=system_instruction)

    def _markdown_to_docx(self, doc, text):
        """
        Simple Markdown parser for Docx.
        Supports: # Headings, - Lists, **Bold**, *Italic*
        """
        for line in text.split('\n'):
            line = line.strip()
            if not line:
                continue
                
            # Headings
            if line.startswith('#'):
                parts = line.split()
                # Count hashes
                level = 0
                for char in parts[0]:
                    if char == '#': level += 1
                    else: break
                
                content = line.lstrip('#').strip()
                # If level > 9, Word cap it, but usually 1-3
                p = doc.add_heading(level=min(level, 3)) # Cap headings in AI text to 3 to avoid giant headers
                self._process_inline_markdown(p, content)
                continue
                
            # List items
            if line.startswith('- ') or line.startswith('* '):
                content = line[2:].strip()
                p = doc.add_paragraph(style='List Bullet')
                self._process_inline_markdown(p, content)
                continue
                
            # Numbered lists (basic 1.)
            if re.match(r'^\d+\.', line):
                content = re.sub(r'^\d+\.', '', line).strip()
                p = doc.add_paragraph(style='List Number')
                self._process_inline_markdown(p, content)
                continue
                
            # Normal paragraph
            p = doc.add_paragraph()
            self._process_inline_markdown(p, line)

    def _process_inline_markdown(self, paragraph, text):
        # Split by bold markers
        # Captures delimiters in result
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**') and len(part) > 3:
                run = paragraph.add_run(part[2:-2])
                run.font.bold = True
            else:
                # Handle italic within non-bold
                # Be careful not to match * inside words unless space? simple *x* is fine
                subparts = re.split(r'(\*.*?\*)', part) 
                for sub in subparts:
                    if sub.startswith('*') and sub.endswith('*') and len(sub) > 2:
                         run = paragraph.add_run(sub[1:-1])
                         run.font.italic = True
                    else:
                         paragraph.add_run(sub)

    def _build_docx(self, scan, stats, grouped_vulns, summary, analyses, conclusion, language):
        doc = Document()
        
        # Severity Translation Map
        sev_map = {
            'CRITICAL': 'CRITICAL', 'HIGH': 'HIGH', 'MEDIUM': 'MEDIUM', 'LOW': 'LOW', 'INFO': 'INFO'
        }
        if language == 'French':
            sev_map = {
                'CRITICAL': 'CRITIQUE', 'HIGH': 'ÉLEVÉ', 'MEDIUM': 'MOYEN', 'LOW': 'FAIBLE', 'INFO': 'INFO'
            }
        elif language == 'Spanish':
             sev_map = {
                'CRITICAL': 'CRÍTICO', 'HIGH': 'ALTO', 'MEDIUM': 'MEDIO', 'LOW': 'BAJO', 'INFO': 'INFO'
            }
        
        # --- TITLE ---
        title = doc.add_heading(f"RAPPORT D'AUDIT DE SÉCURITÉ", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(f"Projet : {scan.project.name}")
        doc.add_paragraph(f"Date : {scan.timestamp.strftime('%d/%m/%Y')}")
        doc.add_paragraph(f"Référence : #SC-{scan.id}")
        doc.add_paragraph(f"Niveau de Confidentialité : INTERNE / CONFIDENTIEL")
        doc.add_paragraph("_" * 50)
        doc.add_page_break()
        
        # --- 1. EXECUTIVE SUMMARY ---
        doc.add_heading('1. SYNTHÈSE EXÉCUTIVE', level=1)
        self._markdown_to_docx(doc, summary)
        
        # Stats Table
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Light Shading Accent 1'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Criticité'
        hdr_cells[1].text = 'Quantité'
        
        for sev in ['CRITICAL', 'HIGH', 'MEDIUM', 'LOW']:
            row = table.add_row().cells
            row[0].text = sev_map.get(sev, sev)
            count = stats.get(sev.lower(), 0)
            row[1].text = str(count)
        
        doc.add_paragraph() # Spacer
        
        # --- 2. SCOPE AND METHODOLOGY ---
        doc.add_heading('2. PÉRIMÈTRE ET MÉTHODOLOGIE', level=1)
        
        doc.add_heading('2.1. Cible de l\'audit', level=2)
        doc.add_paragraph("L'analyse a porté sur l'ensemble du code source et des dépendances des dépôts suivants :")
        
        for repo in scan.project.repositories:
            doc.add_paragraph(f"{repo.name} ({repo.url})", style='List Bullet')
            
        doc.add_paragraph()
        
        doc.add_heading('2.2. Outillage et Standards de Référence', level=2)
        doc.add_paragraph("Cet audit repose sur une approche \"Best-in-Class\", combinant des moteurs d'analyse statique de pointe reconnus pour leur précision.")
        
        doc.add_heading('1. Analyse de la Supply Chain & Infrastructure (SCA/IaC)', level=3)
        doc.add_paragraph("Pour l'analyse des dépendances et de la configuration infrastructure, nous utilisons Trivy, édité par Aqua Security.")
        doc.add_paragraph("SCA : Détection exhaustive des CVEs sur les dépendances.", style='List Bullet')
        doc.add_paragraph("IaC : Audit des configurations Terraform, Docker et Kubernetes (CIS Benchmark).", style='List Bullet')
        
        doc.add_heading('2. Analyse Statique du Code Source (SAST)', level=3)
        doc.add_paragraph("Pour l'analyse de la qualité et de la sécurité du code propriétaire, nous utilisons Semgrep.")
        doc.add_paragraph("OWASP Top 10 (2021) : Couverture complète des risques web critiques.", style='List Bullet')
        doc.add_paragraph("CWE & Secrets : Classification standardisée et détection de secrets hardcodés.", style='List Bullet')

        doc.add_page_break()

        # --- 3. DETAILS ---
        doc.add_heading('3. ANALYSE DÉTAILLÉE DU RISQUE', level=1)
        doc.add_paragraph("Cette section détaille les vulnérabilités les plus critiques identifiées.")
        
        weights = {'CRITICAL': 4, 'HIGH': 3, 'MEDIUM': 2, 'LOW': 1, 'INFO': 0}
        sorted_keys = sorted(
            grouped_vulns.keys(),
            key=lambda k: weights.get(grouped_vulns[k]['severity'], 0),
            reverse=True
        )
        
        for signature in sorted_keys:
            data = grouped_vulns[signature]
            if data['severity'] not in ['CRITICAL', 'HIGH']:
                continue 
                
            # AI Content (Markdown Parsed)
            ai_text = analyses.get(signature, "Analyse automatique non disponible.")
            
            # Extract Translated Title if present
            # PROMPT asked for "TITLE: <Calculated Title>" as first line
            final_title = data['title']
            lines = ai_text.split('\n')
            
            clean_ai_text_lines = []
            
            for line in lines:
                if line.startswith("TITLE:"):
                    # This is the translated title
                    final_title = line.split("TITLE:", 1)[1].strip()
                else:
                    clean_ai_text_lines.append(line)
            
            final_ai_text = "\n".join(clean_ai_text_lines).strip()
            
            # Translate Severity
            translated_severity = sev_map.get(data['severity'], data['severity'])
                
            doc.add_heading(f"{translated_severity} - {final_title}", level=2)
            
            self._markdown_to_docx(doc, final_ai_text)
            
            # Locations
            doc.add_heading("Localisations détectées :", level=3)
            for loc in data['locations'][:10]: 
                p = doc.add_paragraph(style='List Bullet')
                line_str = f":{loc['line']}" if loc['line'] else ""
                text = f"{loc['path']}{line_str}"
                p.add_run(text).font.name = 'Courier New'
                
            if len(data['locations']) > 10:
                doc.add_paragraph(f"... et {len(data['locations']) - 10} autres occurrences.", style='List Bullet')
        
        doc.add_page_break()
        
        # --- 4. INVENTORY TABLE ---
        doc.add_heading('4. INVENTAIRE TECHNIQUE COMPLET', level=1)
        
        inv_table = doc.add_table(rows=1, cols=6)
        inv_table.style = 'Table Grid'
        
        # Headers
        headers = ['ID', 'Sévérité', 'Type', 'Fichier', 'Description', 'Status']
        hdr_row = inv_table.rows[0]
        for idx, text in enumerate(headers):
            cell = hdr_row.cells[idx]
            cell.text = text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

        # Sort all results
        scan_results_sorted = sorted(
            scan.results,
            key=lambda r: weights.get(r.severity, 0),
            reverse=True
        )

        for i, vuln in enumerate(scan_results_sorted, 1):
            row = inv_table.add_row().cells
            row[0].text = f"#{i}"
            row[1].text = vuln.severity
            row[2].text = vuln.owasp_category or "N/A"
            
            path = vuln.file_path
            # No truncation requested by user
            # if len(path) > 30:
            #     path = "..." + path[-27:]
            row[3].text = path
            
            row[4].text = vuln.title
            row[5].text = "À Corriger"
            
        doc.add_paragraph()
        
        # --- 5. CONCLUSION ---
        doc.add_heading('5. CONCLUSION ET PLAN D\'ACTION', level=1)
        self._markdown_to_docx(doc, conclusion)
        
        # --- FOOTER ---
        section = doc.sections[0]
        footer = section.footer
        p = footer.paragraphs[0]
        
        from app.models_settings import Settings
        settings = Settings.query.first()
        company = settings.company_name if settings and settings.company_name else "Spectra"
        
        p.text = f"TLP:AMBER | STRICTEMENT CONFIDENTIEL | © {company} {scan.timestamp.year}"
        
        f = io.BytesIO()
        doc.save(f)
        f.seek(0)
        return f
